from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run


ALIGNMENTS: dict[str, WD_PARAGRAPH_ALIGNMENT] = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
    "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
}

HIGHLIGHT_COLOR_MAP: dict[str, WD_COLOR_INDEX | None] = {
    "yellow": WD_COLOR_INDEX.YELLOW,
    "#ffff00": WD_COLOR_INDEX.YELLOW,
    "ffff00": WD_COLOR_INDEX.YELLOW,
    "green": WD_COLOR_INDEX.BRIGHT_GREEN,
    "brightgreen": WD_COLOR_INDEX.BRIGHT_GREEN,
    "#00ff00": WD_COLOR_INDEX.BRIGHT_GREEN,
    "00ff00": WD_COLOR_INDEX.BRIGHT_GREEN,
    "cyan": WD_COLOR_INDEX.TURQUOISE,
    "#00ffff": WD_COLOR_INDEX.TURQUOISE,
    "00ffff": WD_COLOR_INDEX.TURQUOISE,
    "magenta": WD_COLOR_INDEX.PINK,
    "#ff00ff": WD_COLOR_INDEX.PINK,
    "ff00ff": WD_COLOR_INDEX.PINK,
    "blue": WD_COLOR_INDEX.BLUE,
    "#0000ff": WD_COLOR_INDEX.BLUE,
    "0000ff": WD_COLOR_INDEX.BLUE,
    "red": WD_COLOR_INDEX.RED,
    "#ff0000": WD_COLOR_INDEX.RED,
    "ff0000": WD_COLOR_INDEX.RED,
    "darkblue": WD_COLOR_INDEX.DARK_BLUE,
    "#000080": WD_COLOR_INDEX.DARK_BLUE,
    "000080": WD_COLOR_INDEX.DARK_BLUE,
    "darkcyan": WD_COLOR_INDEX.TEAL,
    "#008080": WD_COLOR_INDEX.TEAL,
    "008080": WD_COLOR_INDEX.TEAL,
    "darkgreen": WD_COLOR_INDEX.GREEN,
    "#008000": WD_COLOR_INDEX.GREEN,
    "008000": WD_COLOR_INDEX.GREEN,
    "darkmagenta": WD_COLOR_INDEX.VIOLET,
    "#800080": WD_COLOR_INDEX.VIOLET,
    "800080": WD_COLOR_INDEX.VIOLET,
    "darkred": WD_COLOR_INDEX.DARK_RED,
    "#800000": WD_COLOR_INDEX.DARK_RED,
    "800000": WD_COLOR_INDEX.DARK_RED,
    "darkyellow": WD_COLOR_INDEX.DARK_YELLOW,
    "#808000": WD_COLOR_INDEX.DARK_YELLOW,
    "808000": WD_COLOR_INDEX.DARK_YELLOW,
    "darkgray": WD_COLOR_INDEX.GRAY_50,
    "#808080": WD_COLOR_INDEX.GRAY_50,
    "808080": WD_COLOR_INDEX.GRAY_50,
    "lightgray": WD_COLOR_INDEX.GRAY_25,
    "#c0c0c0": WD_COLOR_INDEX.GRAY_25,
    "c0c0c0": WD_COLOR_INDEX.GRAY_25,
    "black": WD_COLOR_INDEX.BLACK,
    "#000000": WD_COLOR_INDEX.BLACK,
    "000000": WD_COLOR_INDEX.BLACK,
    "none": None,
}

TOP_LEVEL_RUN_ID_RE = re.compile(r"^p_(?P<paragraph>\d+)_r_(?P<run>\d+)$")
TABLE_RUN_ID_RE = re.compile(
    r"^tbl_(?P<table>\d+)_r_(?P<row>\d+)_c_(?P<cell>\d+)_p_(?P<paragraph>\d+)_r_(?P<run>\d+)$"
)


def _clean_hex_color(raw: Any) -> str | None:
    if not isinstance(raw, str):
        return None
    value = raw.strip().lstrip("#")
    if len(value) != 6:
        return None
    if any(ch not in "0123456789abcdefABCDEF" for ch in value):
        return None
    return value.upper()


def _set_font_name(run: Run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(attr), font_name)


def _normalize_highlight_color(raw: Any) -> WD_COLOR_INDEX | None | object:
    if not isinstance(raw, str):
        return _MISSING
    normalized = raw.strip().replace(" ", "").lower()
    if not normalized:
        return _MISSING
    return HIGHLIGHT_COLOR_MAP.get(normalized, _MISSING)


def _set_run_text(run_element: Any, text: str) -> None:
    for child in list(run_element):
        if child.tag == qn("w:t"):
            run_element.remove(child)
    text_element = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run_element.append(text_element)


_MISSING = object()


def _apply_style(run: Run, paragraph: Paragraph, style: dict[str, Any] | None) -> None:
    if not isinstance(style, dict):
        return

    font_name = style.get("font_name")
    if isinstance(font_name, str) and font_name.strip():
        _set_font_name(run, font_name.strip())

    font_size = style.get("font_size_pt")
    if isinstance(font_size, (int, float)) and float(font_size) > 0:
        run.font.size = Pt(float(font_size))

    is_bold = style.get("is_bold")
    if isinstance(is_bold, bool):
        run.bold = is_bold

    is_italic = style.get("is_italic")
    if isinstance(is_italic, bool):
        run.italic = is_italic

    is_underline = style.get("is_underline")
    if isinstance(is_underline, bool):
        run.underline = is_underline

    is_strike = style.get("is_strike")
    if isinstance(is_strike, bool):
        run.font.strike = is_strike

    is_all_caps = style.get("is_all_caps")
    if isinstance(is_all_caps, bool):
        run.font.all_caps = is_all_caps

    color = _clean_hex_color(style.get("font_color"))
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)

    highlight_color = _normalize_highlight_color(style.get("highlight_color"))
    if highlight_color is not _MISSING:
        run.font.highlight_color = highlight_color

    alignment_name = style.get("paragraph_alignment")
    if isinstance(alignment_name, str):
        alignment = ALIGNMENTS.get(alignment_name.strip().lower())
        if alignment is not None:
            paragraph.alignment = alignment


def _iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_element = parent.element.body
        parent_obj: Any = parent
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
        parent_obj = parent
    else:  # pragma: no cover - defensive
        raise TypeError(f"Unsupported parent type for block iteration: {type(parent)!r}")

    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent_obj)
            continue
        if child.tag == qn("w:tbl"):
            yield Table(child, parent_obj)


def _collect_styles_by_node_id(nodes: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    styles_by_node_id: dict[str, dict[str, Any]] = {}
    for node in nodes:
        if not isinstance(node, dict):
            continue
        node_id = node.get("id")
        style = node.get("style")
        if not isinstance(node_id, str) or not node_id.strip():
            continue
        if not isinstance(style, dict) or not style:
            continue
        styles_by_node_id[node_id.strip()] = style
    return styles_by_node_id


def _apply_styles_to_document(document: DocxDocument, styles_by_node_id: dict[str, dict[str, Any]]) -> None:
    unresolved = set(styles_by_node_id)
    paragraph_idx = 0
    table_idx = 0

    def apply_top_level(parent: DocxDocument | _Cell) -> None:
        nonlocal paragraph_idx, table_idx
        for block in _iter_block_items(parent):
            if isinstance(block, Paragraph):
                apply_paragraph(block, f"p_{paragraph_idx}")
                paragraph_idx += 1
                continue
            if isinstance(block, Table):
                apply_table(block)

    def apply_paragraph(paragraph: Paragraph, paragraph_id: str) -> None:
        for run_idx, run in enumerate(paragraph.runs):
            run_id = f"{paragraph_id}_r_{run_idx}"
            style = styles_by_node_id.get(run_id)
            if style:
                _apply_style(run, paragraph, style)
                unresolved.discard(run_id)

    def apply_table(table: Table) -> None:
        nonlocal table_idx
        current_table_idx = table_idx
        table_idx += 1

        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                paragraph_local_idx = 0
                for block in _iter_block_items(cell):
                    if isinstance(block, Paragraph):
                        paragraph_id = (
                            f"tbl_{current_table_idx}_r_{row_idx}_c_{cell_idx}_p_{paragraph_local_idx}"
                        )
                        apply_paragraph(block, paragraph_id)
                        paragraph_local_idx += 1
                        continue
                    if isinstance(block, Table):
                        apply_table(block)

    apply_top_level(document)

    if unresolved:
        unresolved_list = ", ".join(sorted(unresolved)[:10])
        raise ValueError(f"Unable to map style updates to source DOCX nodes: {unresolved_list}")


def _write_rebuilt_docx(nodes: list[dict[str, Any]], output_docx: Path) -> None:
    document = Document()

    for node in nodes:
        if not isinstance(node, dict):
            continue
        text = node.get("text", "")
        if not isinstance(text, str):
            text = str(text)
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        _apply_style(run, paragraph, node.get("style") if isinstance(node.get("style"), dict) else None)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))


def _find_table_by_index(parent: DocxDocument | _Cell, target_index: int) -> Table:
    current_index = 0

    def visit(container: DocxDocument | _Cell) -> Table | None:
        nonlocal current_index
        for block in _iter_block_items(container):
            if isinstance(block, Paragraph):
                continue
            if current_index == target_index:
                return block
            current_index += 1
            for row in block.rows:
                for cell in row.cells:
                    nested = visit(cell)
                    if nested is not None:
                        return nested
        return None

    table = visit(parent)
    if table is None:
        raise ValueError(f"Unable to resolve table index {target_index}.")
    return table


def _locate_run(document: DocxDocument, node_id: str) -> tuple[Paragraph, int]:
    match = TOP_LEVEL_RUN_ID_RE.fullmatch(node_id)
    if match:
        paragraph = document.paragraphs[int(match.group("paragraph"))]
        run_index = int(match.group("run"))
        if run_index >= len(paragraph.runs):
            raise ValueError(f"Unable to resolve run index {run_index} for node {node_id}.")
        return paragraph, run_index

    match = TABLE_RUN_ID_RE.fullmatch(node_id)
    if not match:
        raise ValueError(f"Unsupported node id for structural operation: {node_id}")

    table = _find_table_by_index(document, int(match.group("table")))
    row = table.rows[int(match.group("row"))]
    cell = row.cells[int(match.group("cell"))]
    paragraph = cell.paragraphs[int(match.group("paragraph"))]
    run_index = int(match.group("run"))
    if run_index >= len(paragraph.runs):
        raise ValueError(f"Unable to resolve run index {run_index} for node {node_id}.")
    return paragraph, run_index


def _merge_paragraph(document: DocxDocument, target_node_id: str) -> None:
    paragraph, _ = _locate_run(document, target_node_id)
    parent = paragraph._element.getparent()
    siblings = list(parent)
    paragraph_index = siblings.index(paragraph._element)
    next_paragraph = None
    for sibling in siblings[paragraph_index + 1 :]:
        if sibling.tag == qn("w:p"):
            next_paragraph = sibling
            break
    if next_paragraph is None:
        raise ValueError(f"merge_paragraph requires a following paragraph: {target_node_id}")

    for child in list(next_paragraph):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._element.append(deepcopy(child))
    parent.remove(next_paragraph)


def _split_paragraph(document: DocxDocument, target_node_id: str, split_offset: int) -> None:
    paragraph, run_index = _locate_run(document, target_node_id)
    run = paragraph.runs[run_index]
    text = run.text or ""
    if split_offset <= 0 or split_offset >= len(text):
        raise ValueError(f"split_paragraph requires split_offset inside run text: {target_node_id}")

    prefix = text[:split_offset]
    suffix = text[split_offset:]
    run.text = prefix

    source_paragraph = paragraph._element
    new_paragraph = OxmlElement("w:p")
    if source_paragraph.pPr is not None:
        new_paragraph.append(deepcopy(source_paragraph.pPr))
    source_paragraph.addnext(new_paragraph)

    suffix_run = deepcopy(run._element)
    _set_run_text(suffix_run, suffix)
    new_paragraph.append(suffix_run)

    found_current_run = False
    for child in list(source_paragraph):
        if child is run._element:
            found_current_run = True
            continue
        if not found_current_run:
            continue
        source_paragraph.remove(child)
        new_paragraph.append(child)


def apply_structure_operation(input_docx: Path, operation: dict[str, Any]) -> None:
    document = Document(str(input_docx))
    target_node_id = str(operation.get("targetNodeId", "")).strip()
    operation_type = str(operation.get("type", "")).strip()
    payload = operation.get("payload")
    if not isinstance(payload, dict):
        payload = {}

    if operation_type == "merge_paragraph":
        _merge_paragraph(document, target_node_id)
    elif operation_type == "split_paragraph":
        split_offset = payload.get("split_offset")
        if not isinstance(split_offset, int):
            raise ValueError("split_paragraph requires payload.split_offset")
        _split_paragraph(document, target_node_id, split_offset)
    else:
        raise ValueError(f"Unsupported structure operation: {operation_type}")

    document.save(str(input_docx))


def write_docx_from_ir(input_json: Path, output_docx: Path) -> None:
    payload = json.loads(input_json.read_text(encoding="utf-8"))
    nodes = payload.get("nodes", [])
    if not isinstance(nodes, list):
        raise ValueError("DocumentIR.nodes must be a list")

    metadata = payload.get("metadata")
    input_docx_path = None
    if isinstance(metadata, dict):
        raw_path = metadata.get("workingDocxPath") or metadata.get("inputDocxPath")
        if isinstance(raw_path, str) and raw_path.strip():
            input_docx_path = Path(raw_path.strip())

    output_docx.parent.mkdir(parents=True, exist_ok=True)

    if input_docx_path is None:
        _write_rebuilt_docx(nodes, output_docx)
        return

    document = Document(str(input_docx_path))
    _apply_styles_to_document(document, _collect_styles_by_node_id(nodes))
    document.save(str(output_docx))


def main() -> int:
    parser = argparse.ArgumentParser(description="Write DocumentIR JSON to .docx")
    parser.add_argument("--input-json", required=True, help="Path to DocumentIR JSON file")
    parser.add_argument("--output-docx", required=True, help="Target .docx file path")
    args = parser.parse_args()

    write_docx_from_ir(Path(args.input_json), Path(args.output_docx))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
