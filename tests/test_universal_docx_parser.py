from __future__ import annotations

import base64
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch

import src.python.core.universal_docx_parser as parser_module
from src.python.core.universal_docx_parser import UniversalDocxParser


_PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+X2fQAAAAASUVORK5CYII="
)


class UniversalDocxParserTest(unittest.TestCase):
    def test_module_defines_runtime_docxdocument_symbol(self) -> None:
        self.assertTrue(hasattr(parser_module, "DocxDocument"))

    def test_extract_doc_defaults_supports_baseoxml_elements_without_namespaces_kw(self) -> None:
        parser = UniversalDocxParser.__new__(UniversalDocxParser)
        styles_root = _FakeOxmlElement(
            "w:styles",
            children=[
                _FakeOxmlElement(
                    "w:docDefaults",
                    children=[
                        _FakeOxmlElement(
                            "w:rPrDefault",
                            children=[
                                _FakeOxmlElement(
                                    "w:rPr",
                                    children=[
                                        _FakeOxmlElement(
                                            "w:rFonts",
                                            attrib={"w:eastAsia": "SimSun", "w:ascii": "Calibri"},
                                        ),
                                        _FakeOxmlElement("w:sz", attrib={"w:val": "21"}),
                                        _FakeOxmlElement("w:color", attrib={"w:val": "112233"}),
                                    ],
                                )
                            ],
                        )
                    ],
                )
            ],
        )
        fake_document = SimpleNamespace(styles=SimpleNamespace(element=styles_root))

        with patch.object(parser_module, "qn", side_effect=lambda name: name):
            defaults = parser._extract_doc_defaults(fake_document)

        self.assertEqual(defaults.font_name, "SimSun")
        self.assertEqual(defaults.font_size_pt, 10.5)
        self.assertEqual(defaults.font_color, "112233")

    def test_extract_image_node_supports_baseoxml_elements_without_namespaces_kw(self) -> None:
        parser = UniversalDocxParser.__new__(UniversalDocxParser)
        parser._image_idx = 0

        drawing = _FakeOxmlElement(
            "w:drawing",
            children=[
                _FakeOxmlElement(
                    "wp:inline",
                    children=[
                        _FakeOxmlElement("wp:extent", attrib={"cx": "9525", "cy": "19050"}),
                        _FakeOxmlElement(
                            "a:graphic",
                            children=[_FakeOxmlElement("a:blip", attrib={"r:embed": "rId5"})],
                        ),
                    ],
                )
            ],
        )

        with patch.object(parser_module, "qn", side_effect=lambda name: name):
            with patch.object(parser, "_export_image_by_rid", return_value={"id": "img_0"}) as export:
                result = parser._extract_image_node(object(), drawing)

        self.assertEqual(result, {"id": "img_0"})
        export.assert_called_once()
        self.assertEqual(export.call_args.args[1], "rId5")
        self.assertEqual(export.call_args.args[2], "img_0")
        self.assertEqual(export.call_args.args[3], {"width": 1.0, "height": 2.0})

    def test_extract_vml_images_supports_baseoxml_elements_without_namespaces_kw(self) -> None:
        parser = UniversalDocxParser.__new__(UniversalDocxParser)
        parser._image_idx = 0

        pict = _FakeOxmlElement(
            "w:pict",
            children=[
                _FakeOxmlElement(
                    "v:shape",
                    attrib={"style": "width:12pt;height:24px"},
                    children=[
                        _FakeOxmlElement(
                            parser_module.UniversalDocxParser._ns_tag("v", "imagedata"),
                            attrib={"r:id": "rId7"},
                        )
                    ],
                )
            ],
        )

        with patch.object(parser_module, "qn", side_effect=lambda name: name):
            with patch.object(parser, "_export_image_by_rid", return_value={"id": "img_0"}) as export:
                nodes = parser._extract_vml_images(object(), pict)

        self.assertEqual(nodes, [{"id": "img_0"}])
        export.assert_called_once()
        self.assertEqual(export.call_args.args[1], "rId7")
        self.assertEqual(export.call_args.args[2], "img_0")
        self.assertEqual(export.call_args.args[3], {"width": 16.0, "height": 24.0})

    def test_parse_docx_with_text_table_image_and_formula(self) -> None:
        docx = self._import_docx()
        parser_mod = self._import_docx_parser_xml()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            media_root = tmp_path / "agent_workspace" / "media"
            media_root.mkdir(parents=True, exist_ok=True)

            image_path = tmp_path / "test.png"
            image_path.write_bytes(_PNG_1X1)

            doc = docx.Document()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("测试文本")
            run.bold = True
            run.italic = False
            run.font.name = "SimSun"
            run.font.size = docx.shared.Pt(12)
            run.add_picture(str(image_path))

            formula_run = paragraph.add_run()
            formula_xml = (
                '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
                "<m:sSup><m:e><m:r><m:t>x</m:t></m:r></m:e>"
                "<m:sup><m:r><m:t>2</m:t></m:r></m:sup></m:sSup>"
                "</m:oMath>"
            )
            formula_run._r.append(parser_mod.parse_xml(formula_xml))

            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "表格内容"

            docx_path = tmp_path / "sample.docx"
            doc.save(str(docx_path))

            parser = UniversalDocxParser(media_dir=media_root)
            state = parser.parse(docx_path)

            self.assertIn("document_meta", state)
            self.assertIn("nodes", state)
            self.assertEqual(state["document_meta"]["total_tables"], 1)

            paragraph_node = state["nodes"][0]
            self.assertEqual(paragraph_node["id"], "p_0")
            self.assertEqual(paragraph_node["node_type"], "paragraph")
            self.assertTrue(paragraph_node["children"])

            text_nodes = [n for n in paragraph_node["children"] if n["node_type"] == "text_run"]
            self.assertTrue(text_nodes)
            style = text_nodes[0]["style"]
            self.assertIsInstance(style["font_name"], str)
            self.assertIsInstance(style["font_size_pt"], float)
            self.assertIsInstance(style["font_color"], str)
            self.assertIsInstance(style["is_bold"], bool)
            self.assertIsInstance(style["is_italic"], bool)

            image_nodes = [n for n in paragraph_node["children"] if n["node_type"] == "image"]
            self.assertTrue(image_nodes)
            self.assertTrue(str(image_nodes[0]["id"]).startswith("img_"))
            self.assertTrue(Path(image_nodes[0]["src"]).is_absolute())
            self.assertTrue(Path(image_nodes[0]["src"]).exists())
            self.assertIn("width", image_nodes[0]["size"])
            self.assertIn("height", image_nodes[0]["size"])

            formula_nodes = [n for n in paragraph_node["children"] if n["node_type"] == "formula"]
            self.assertTrue(formula_nodes)
            self.assertEqual(formula_nodes[0]["format"], "latex")
            self.assertIn("x", formula_nodes[0]["content"])

            table_node = [n for n in state["nodes"] if n["node_type"] == "table"][0]
            self.assertEqual(table_node["id"], "tbl_0")
            run_node = table_node["rows"][0]["cells"][0]["paragraphs"][0]["children"][0]
            self.assertEqual(run_node["id"], "tbl_0_r_0_c_0_p_0_r_0")
            self.assertEqual(run_node["node_type"], "text_run")
            self.assertEqual(run_node["content"], "表格内容")

    def test_classify_paragraph_role_prefers_heading_then_list_then_body(self) -> None:
        parser = UniversalDocxParser.__new__(UniversalDocxParser)

        heading_role = parser._classify_paragraph_role(
            _FakeParagraph(style_name="Heading 2", text="章节标题"),
            in_table=False,
        )
        list_role = parser._classify_paragraph_role(
            _FakeParagraph(style_name="List Paragraph", text="1. 列表", list_level=0),
            in_table=False,
        )
        body_role = parser._classify_paragraph_role(
            _FakeParagraph(style_name="Normal", text="普通正文"),
            in_table=False,
        )

        self.assertEqual(heading_role["role"], "heading")
        self.assertEqual(heading_role["heading_level"], 2)
        self.assertEqual(list_role["role"], "list_item")
        self.assertEqual(list_role["list_level"], 0)
        self.assertEqual(body_role["role"], "body")

    @staticmethod
    def _import_docx():
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc
        return docx

    @staticmethod
    def _import_docx_parser_xml():
        try:
            from docx.oxml import parse_xml  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"docx xml parser not available: {exc}") from exc
        return type("DocxXmlTools", (), {"parse_xml": parse_xml})

class _FakeOxmlElement:
    def __init__(
        self,
        tag: str,
        attrib: dict[str, str] | None = None,
        children: list["_FakeOxmlElement"] | None = None,
    ) -> None:
        self.tag = tag
        self._attrib = attrib or {}
        self._children = children or []
        self._parent: "_FakeOxmlElement | None" = None
        for child in self._children:
            child._parent = self

    def get(self, key: str, default: str | None = None) -> str | None:
        return self._attrib.get(key, default)

    def getparent(self) -> "_FakeOxmlElement | None":
        return self._parent

    def iterchildren(self):
        return iter(self._children)

    def __iter__(self):
        return iter(self._children)

    def xpath(self, expr: str):
        return []


class _FakeParagraphStyle:
    def __init__(self, name: str) -> None:
        self.name = name


class _FakeParagraphElement:
    def __init__(self, list_level: int | None) -> None:
        self.list_level = list_level


class _FakeParagraph:
    def __init__(self, style_name: str, text: str, list_level: int | None = None) -> None:
        self.style = _FakeParagraphStyle(style_name)
        self.text = text
        self._p = _FakeParagraphElement(list_level)


if __name__ == "__main__":
    unittest.main()
