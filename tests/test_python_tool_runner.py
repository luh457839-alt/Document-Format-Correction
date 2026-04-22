from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from src.python.api.python_tool_runner import PythonToolRunnerError, execute_tool_request


class PythonToolRunnerTest(unittest.TestCase):
    def test_inspect_document_returns_summary_without_mutating_doc(self) -> None:
        source_doc = {
            "id": "doc1",
            "version": "v1",
            "nodes": [{"id": "n1", "text": "hello"}],
            "metadata": {"keep": True},
        }
        result = execute_tool_request(
            {
                "action": "execute",
                "toolName": "inspect_document",
                "input": {
                    "doc": source_doc,
                    "context": {"taskId": "t1", "stepId": "s1", "dryRun": False},
                },
            }
        )

        self.assertEqual(result["summary"], "Inspected 1 node(s).")
        self.assertEqual(result["doc"], source_doc)
        self.assertIsNot(result["doc"], source_doc)

    def test_docx_observation_attaches_observation_to_metadata(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            docx_path = tmp_path / "sample.docx"
            document = docx.Document()
            document.add_paragraph("hello")
            document.save(docx_path)

            result = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "docx_observation",
                    "input": {
                        "doc": {"id": "doc1", "version": "v1", "nodes": []},
                        "operation": {
                            "id": "op1",
                            "type": "set_font",
                            "targetNodeId": "n1",
                            "payload": {"docxPath": str(docx_path)},
                        },
                        "context": {"taskId": "t1", "stepId": "s1", "dryRun": False},
                    },
                }
            )

            observation = result["doc"]["metadata"]["docxObservation"]
            self.assertIn("document_meta", observation)
            self.assertGreaterEqual(observation["document_meta"]["total_paragraphs"], 1)

    def test_write_operation_only_updates_doc_until_materialize(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            document = docx.Document()
            paragraph = document.add_paragraph()
            run = paragraph.add_run("hello")
            run.font.name = "Arial"
            document.save(source)

            request = {
                "action": "execute",
                "toolName": "write_operation",
                "input": {
                    "doc": {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [{"id": "p_0_r_0", "text": "hello"}],
                        "metadata": {
                            "inputDocxPath": str(source),
                            "outputDocxPath": str(output),
                        },
                    },
                    "operation": {
                        "id": "op1",
                        "type": "set_font",
                        "targetNodeId": "p_0_r_0",
                        "payload": {"font_name": "SimSun"},
                    },
                    "context": {"taskId": "t1", "stepId": "s1", "dryRun": False},
                },
            }

            executed = execute_tool_request(request)
            self.assertFalse(output.exists())
            self.assertEqual(executed["summary"], "Applied set_font to p_0_r_0.")
            self.assertEqual(
                executed["doc"]["nodes"][0]["style"]["font_name"],
                "SimSun",
            )

            materialized = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": executed["doc"],
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )
            self.assertTrue(output.exists())
            self.assertEqual(materialized["artifacts"]["outputDocxPath"], str(output))

    def test_write_operation_batches_multiple_nodes_in_one_request(self) -> None:
        executed = execute_tool_request(
            {
                "action": "execute",
                "toolName": "write_operation",
                "input": {
                    "doc": {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {"id": "p_0_r_0", "text": "标题"},
                            {"id": "p_1_r_0", "text": "第一段"},
                            {"id": "p_1_r_1", "text": "正文"},
                        ],
                    },
                    "operation": {
                        "id": "op_batch",
                        "type": "set_font",
                        "targetNodeIds": ["p_0_r_0", "p_1_r_0", "p_1_r_1"],
                        "payload": {"font_name": "SimSun"},
                    },
                    "context": {"taskId": "t1", "stepId": "s_batch", "dryRun": False},
                },
            }
        )

        self.assertEqual(executed["summary"], "Applied set_font to 3 nodes.")
        self.assertEqual(
            [node.get("style", {}).get("font_name") for node in executed["doc"]["nodes"]],
            ["SimSun", "SimSun", "SimSun"],
        )

    def test_write_operation_normalizes_exact_line_spacing(self) -> None:
        executed = execute_tool_request(
            {
                "action": "execute",
                "toolName": "write_operation",
                "input": {
                    "doc": {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [{"id": "p_0_r_0", "text": "正文"}],
                        "metadata": {"outputDocxPath": "output.docx"},
                    },
                    "operation": {
                        "id": "op_line_spacing",
                        "type": "set_line_spacing",
                        "targetNodeId": "p_0_r_0",
                        "payload": {"line_spacing": {"mode": "exact", "pt": 20}},
                    },
                    "context": {"taskId": "t1", "stepId": "s_line_spacing", "dryRun": False},
                },
            }
        )

        self.assertEqual(
            executed["doc"]["nodes"][0]["style"]["line_spacing"],
            {"mode": "exact", "pt": 20.0},
        )
        self.assertEqual(
            executed["summary"],
            "Applied set_line_spacing to p_0_r_0; pending materialize to output.docx.",
        )

    def test_write_operation_rejects_invalid_line_spacing(self) -> None:
        with self.assertRaises(PythonToolRunnerError) as ctx:
            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": {
                            "id": "doc1",
                            "version": "v1",
                            "nodes": [{"id": "p_0_r_0", "text": "正文"}],
                        },
                        "operation": {
                            "id": "op_line_spacing",
                            "type": "set_line_spacing",
                            "targetNodeId": "p_0_r_0",
                            "payload": {"line_spacing": {"mode": "exact"}},
                        },
                        "context": {"taskId": "t1", "stepId": "s_line_spacing", "dryRun": False},
                    },
                }
            )

        self.assertEqual(ctx.exception.code, "E_INVALID_OPERATION_PAYLOAD")

    def test_materialize_document_writes_output_once(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            document = docx.Document()
            paragraph = document.add_paragraph()
            run = paragraph.add_run("hello")
            run.font.name = "Arial"
            document.save(source)

            result = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": {
                            "id": "doc1",
                            "version": "v1",
                            "nodes": [
                                {
                                    "id": "p_0_r_0",
                                    "text": "hello",
                                    "style": {"font_name": "SimSun"},
                                }
                            ],
                            "metadata": {
                                "inputDocxPath": str(source),
                                "outputDocxPath": str(output),
                            },
                        },
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )

            self.assertTrue(output.exists())
            self.assertEqual(
                result["artifacts"]["outputDocxPath"],
                str(output),
            )

    def test_materialize_document_writes_line_spacing_when_source_had_no_explicit_spacing(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            document.save(source)

            observation = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "docx_observation",
                    "input": {
                        "doc": {"id": "doc1", "version": "v1", "nodes": []},
                        "operation": {
                            "id": "observe",
                            "type": "set_font",
                            "targetNodeId": "p_0_r_0",
                            "payload": {"docxPath": str(source)},
                        },
                        "context": {"taskId": "t1", "stepId": "observe", "dryRun": False},
                    },
                }
            )
            observed_style = observation["doc"]["metadata"]["docxObservation"]["nodes"][0]["children"][0]["style"]
            self.assertNotIn("line_spacing", observed_style)

            executed = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": {
                            "id": "doc1",
                            "version": "v1",
                            "nodes": [{"id": "p_0_r_0", "text": "hello"}],
                            "metadata": {
                                "inputDocxPath": str(source),
                                "outputDocxPath": str(output),
                            },
                        },
                        "operation": {
                            "id": "set_line_spacing",
                            "type": "set_line_spacing",
                            "targetNodeId": "p_0_r_0",
                            "payload": {"line_spacing": 1.5},
                        },
                        "context": {"taskId": "t1", "stepId": "set_line_spacing", "dryRun": False},
                    },
                }
            )

            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": executed["doc"],
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )

            result = docx.Document(output)
            self.assertEqual(result.paragraphs[0].paragraph_format.line_spacing, 1.5)

    def test_materialize_document_rejects_docx_mapped_nodes_without_source_path(self) -> None:
        with self.assertRaises(PythonToolRunnerError) as ctx:
            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": {
                            "id": "doc1",
                            "version": "v1",
                            "nodes": [
                                {
                                    "id": "p_0_r_0",
                                    "text": "hello",
                                    "style": {"line_spacing": 1.5},
                                }
                            ],
                            "metadata": {
                                "outputDocxPath": "output.docx",
                                "sourceDocumentMeta": {"total_paragraphs": 1},
                            },
                        },
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )

        self.assertEqual(ctx.exception.code, "E_INPUT_PATH_REQUIRED")
        self.assertIn("inputDocxPath", ctx.exception.message)

    def test_line_spacing_materialize_preserves_alignment_and_size_with_multiple_runs(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            paragraph.add_run("world")
            document.save(source)

            doc = {
                "id": "doc1",
                "version": "v1",
                "nodes": [
                    {"id": "p_0_r_0", "text": "hello"},
                    {"id": "p_0_r_1", "text": "world"},
                ],
                "metadata": {
                    "inputDocxPath": str(source),
                    "outputDocxPath": str(output),
                },
            }

            sized = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": doc,
                        "operation": {
                            "id": "set_size",
                            "type": "set_size",
                            "targetNodeId": "p_0_r_0",
                            "payload": {"font_size_pt": 16},
                        },
                        "context": {"taskId": "t1", "stepId": "set_size", "dryRun": False},
                    },
                }
            )["doc"]
            aligned = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": sized,
                        "operation": {
                            "id": "set_alignment",
                            "type": "set_alignment",
                            "targetNodeId": "p_0_r_0",
                            "payload": {"paragraph_alignment": "center"},
                        },
                        "context": {"taskId": "t1", "stepId": "set_alignment", "dryRun": False},
                    },
                }
            )["doc"]
            spaced = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": aligned,
                        "operation": {
                            "id": "set_line_spacing",
                            "type": "set_line_spacing",
                            "targetNodeIds": ["p_0_r_0", "p_0_r_1"],
                            "payload": {"line_spacing": {"mode": "exact", "pt": 18}},
                        },
                        "context": {"taskId": "t1", "stepId": "set_line_spacing", "dryRun": False},
                    },
                }
            )["doc"]

            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": spaced,
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )

            result = docx.Document(output)
            paragraph = result.paragraphs[0]
            self.assertEqual(paragraph.alignment, docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER)
            self.assertEqual(paragraph.paragraph_format.line_spacing.pt, 18.0)
            self.assertEqual(paragraph.runs[0].font.size.pt, 16.0)

    def test_merge_paragraph_rehydrates_nodes_before_materialize(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            document = docx.Document()
            document.add_paragraph("hello")
            document.add_paragraph("world")
            document.save(source)

            merged = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": {
                            "id": "doc1",
                            "version": "v1",
                            "nodes": [
                                {"id": "p_0_r_0", "text": "hello"},
                                {"id": "p_1_r_0", "text": "world"},
                            ],
                            "metadata": {
                                "inputDocxPath": str(source),
                                "outputDocxPath": str(output),
                            },
                        },
                        "operation": {
                            "id": "op_merge",
                            "type": "merge_paragraph",
                            "targetNodeId": "p_0_r_0",
                            "payload": {},
                        },
                        "context": {"taskId": "t1", "stepId": "s_merge", "dryRun": False},
                    },
                }
            )

            self.assertEqual([node["id"] for node in merged["doc"]["nodes"]], ["p_0_r_0", "p_0_r_1"])
            self.assertIn("workingDocxPath", merged["doc"]["metadata"])

            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": merged["doc"],
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )

            result = docx.Document(output)
            self.assertEqual(len(result.paragraphs), 1)
            self.assertEqual(result.paragraphs[0].text, "helloworld")

    def test_split_paragraph_rehydrates_nodes_before_materialize(self) -> None:
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            document = docx.Document()
            document.add_paragraph("helloworld")
            document.save(source)

            split = execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "write_operation",
                    "input": {
                        "doc": {
                            "id": "doc1",
                            "version": "v1",
                            "nodes": [{"id": "p_0_r_0", "text": "helloworld"}],
                            "metadata": {
                                "inputDocxPath": str(source),
                                "outputDocxPath": str(output),
                            },
                        },
                        "operation": {
                            "id": "op_split",
                            "type": "split_paragraph",
                            "targetNodeId": "p_0_r_0",
                            "payload": {"split_offset": 5},
                        },
                        "context": {"taskId": "t1", "stepId": "s_split", "dryRun": False},
                    },
                }
            )

            self.assertEqual([node["id"] for node in split["doc"]["nodes"]], ["p_0_r_0", "p_1_r_0"])

            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "materialize_document",
                    "input": {
                        "doc": split["doc"],
                        "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                    },
                }
            )

            result = docx.Document(output)
            self.assertEqual([paragraph.text for paragraph in result.paragraphs], ["hello", "world"])

    def test_unknown_tool_raises_structured_error(self) -> None:
        with self.assertRaises(PythonToolRunnerError) as ctx:
            execute_tool_request(
                {
                    "action": "execute",
                    "toolName": "unknown_tool",
                    "input": {
                        "doc": {"id": "doc1", "version": "v1", "nodes": []},
                        "context": {"taskId": "t1", "stepId": "s1", "dryRun": False},
                    },
                }
            )

        self.assertEqual(ctx.exception.code, "E_TOOL_NOT_FOUND")

    def test_docx_observation_dependency_failures_are_classified(self) -> None:
        with patch(
            "src.python.api.python_tool_runner._parse_docx_state",
            side_effect=RuntimeError("lxml is required for UniversalDocxParser: import failed"),
        ):
            with self.assertRaises(PythonToolRunnerError) as ctx:
                execute_tool_request(
                    {
                        "action": "execute",
                        "toolName": "docx_observation",
                        "input": {
                            "doc": {"id": "doc1", "version": "v1", "nodes": []},
                            "operation": {
                                "id": "op1",
                                "type": "set_font",
                                "targetNodeId": "n1",
                                "payload": {"docxPath": "missing.docx"},
                            },
                            "context": {"taskId": "t1", "stepId": "s1", "dryRun": False},
                        },
                    }
                )

        self.assertEqual(ctx.exception.code, "E_PYTHON_DEPENDENCY_MISSING")

    def test_materialize_dependency_failures_are_classified(self) -> None:
        with patch(
            "src.python.api.python_tool_runner._load_docx_writer_functions",
            side_effect=ModuleNotFoundError("No module named 'docx'"),
        ):
            with self.assertRaises(PythonToolRunnerError) as ctx:
                execute_tool_request(
                    {
                        "action": "execute",
                        "toolName": "materialize_document",
                        "input": {
                            "doc": {
                                "id": "doc1",
                                "version": "v1",
                                "nodes": [{"id": "p_0_r_0", "text": "hello"}],
                                "metadata": {
                                    "inputDocxPath": "input.docx",
                                    "outputDocxPath": "output.docx",
                                },
                            },
                            "context": {"taskId": "t1", "stepId": "finalize", "dryRun": False},
                        },
                    }
                )

        self.assertEqual(ctx.exception.code, "E_PYTHON_DEPENDENCY_MISSING")


if __name__ == "__main__":
    unittest.main()
