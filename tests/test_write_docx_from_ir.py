from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path


class WriteDocxFromIrTest(unittest.TestCase):
    def test_in_place_size_change_preserves_existing_font_name(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            run = paragraph.add_run("hello")
            run.font.name = "Arial"
            run.font.size = docx.shared.Pt(11)
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"font_size_pt": 22},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            written_run = result.paragraphs[0].runs[0]
            self.assertEqual(written_run.font.size.pt, 22.0)
            self.assertEqual(written_run.font.name, "Arial")

    def test_in_place_font_change_preserves_existing_size(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            run = paragraph.add_run("hello")
            run.font.name = "Arial"
            run.font.size = docx.shared.Pt(11)
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"font_name": "SimSun"},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            written_run = result.paragraphs[0].runs[0]
            self.assertEqual(written_run.font.name, "SimSun")
            self.assertEqual(written_run.font.size.pt, 11.0)

    def test_in_place_extended_style_fields_are_written(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {
                                    "font_color": "112233",
                                    "is_bold": True,
                                    "is_italic": True,
                                    "is_underline": True,
                                    "is_strike": True,
                                    "highlight_color": "yellow",
                                    "is_all_caps": True,
                                    "paragraph_alignment": "center",
                                },
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            written_paragraph = result.paragraphs[0]
            written_run = written_paragraph.runs[0]
            self.assertEqual(str(written_run.font.color.rgb), "112233")
            self.assertTrue(written_run.bold)
            self.assertTrue(written_run.italic)
            self.assertTrue(written_run.underline)
            self.assertTrue(written_run.font.strike)
            self.assertEqual(written_run.font.highlight_color, docx.enum.text.WD_COLOR_INDEX.YELLOW)
            self.assertTrue(written_run.font.all_caps)
            self.assertEqual(written_paragraph.alignment, docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER)

    def test_in_place_line_spacing_fields_are_written(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output_multiple = tmp_path / "output-multiple.docx"
            output_exact = tmp_path / "output-exact.docx"
            payload_multiple_path = tmp_path / "payload-multiple.json"
            payload_exact_path = tmp_path / "payload-exact.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            document.save(source)

            payload_multiple_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"line_spacing": 1.5},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            payload_exact_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"line_spacing": {"mode": "exact", "pt": 20}},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_multiple_path, output_multiple)
            write_docx_from_ir(payload_exact_path, output_exact)

            multiple = docx.Document(output_multiple)
            exact = docx.Document(output_exact)
            self.assertEqual(multiple.paragraphs[0].paragraph_format.line_spacing, 1.5)
            self.assertEqual(exact.paragraphs[0].paragraph_format.line_spacing.pt, 20.0)

    def test_page_spacing_and_first_line_indent_fields_are_written(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "n1",
                                "text": "hello",
                                "style": {
                                    "space_before_pt": 6,
                                    "space_after_pt": 3,
                                    "first_line_indent_pt": 24,
                                },
                            }
                        ],
                        "metadata": {
                            "page_layout": {
                                "paper_size": "A4",
                                "margin_top_cm": 3.7,
                                "margin_bottom_cm": 3.5,
                                "margin_left_cm": 2.8,
                                "margin_right_cm": 2.6,
                            }
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            section = result.sections[0]
            paragraph_format = result.paragraphs[0].paragraph_format
            self.assertAlmostEqual(section.page_width.cm, 21.0, places=1)
            self.assertAlmostEqual(section.page_height.cm, 29.7, places=1)
            self.assertAlmostEqual(section.top_margin.cm, 3.7, places=1)
            self.assertAlmostEqual(section.bottom_margin.cm, 3.5, places=1)
            self.assertAlmostEqual(section.left_margin.cm, 2.8, places=1)
            self.assertAlmostEqual(section.right_margin.cm, 2.6, places=1)
            self.assertEqual(paragraph_format.space_before.pt, 6.0)
            self.assertEqual(paragraph_format.space_after.pt, 3.0)
            self.assertEqual(paragraph_format.first_line_indent.pt, 24.0)

    def test_in_place_write_requires_source_docx_for_docx_mapped_nodes(self) -> None:
        self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"line_spacing": 1.5},
                            }
                        ],
                        "metadata": {
                            "sourceDocumentMeta": {"total_paragraphs": 1},
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            with self.assertRaisesRegex(ValueError, "inputDocxPath"):
                write_docx_from_ir(payload_path, output)

    def test_in_place_hex_highlight_alias_is_supported(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"highlight_color": "#FFFF00"},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            written_run = result.paragraphs[0].runs[0]
            self.assertEqual(written_run.font.highlight_color, docx.enum.text.WD_COLOR_INDEX.YELLOW)

    def test_in_place_bright_green_highlight_alias_is_supported(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"highlight_color": "#00FF00"},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            written_run = result.paragraphs[0].runs[0]
            self.assertEqual(written_run.font.highlight_color, docx.enum.text.WD_COLOR_INDEX.BRIGHT_GREEN)

    def test_in_place_dark_green_highlight_alias_is_supported(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("hello")
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0",
                                "text": "hello",
                                "style": {"highlight_color": "#008000"},
                            }
                        ],
                        "metadata": {"inputDocxPath": str(source)},
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            written_run = result.paragraphs[0].runs[0]
            self.assertEqual(written_run.font.highlight_color, docx.enum.text.WD_COLOR_INDEX.GREEN)

    def test_rebuilds_split_runs_from_source_run_ids(self) -> None:
        docx = self._import_docx()
        write_docx_from_ir = self._import_writer()

        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_path = Path(tmp_dir)
            source = tmp_path / "source.docx"
            output = tmp_path / "output.docx"
            payload_path = tmp_path / "payload.json"

            document = docx.Document()
            paragraph = document.add_paragraph()
            run = paragraph.add_run("这是NLP2025报告")
            run.font.name = "FangSong_GB2312"
            run.font.size = docx.shared.Pt(16)
            run.bold = True
            document.save(source)

            payload_path.write_text(
                json.dumps(
                    {
                        "id": "doc1",
                        "version": "v1",
                        "nodes": [
                            {
                                "id": "p_0_r_0__seg_0",
                                "sourceRunId": "p_0_r_0",
                                "text": "这是",
                                "style": {
                                    "font_name": "FangSong_GB2312",
                                    "font_size_pt": 16,
                                    "is_bold": True,
                                },
                            },
                            {
                                "id": "p_0_r_0__seg_1",
                                "sourceRunId": "p_0_r_0",
                                "text": "NLP2025",
                                "style": {
                                    "font_name": "Times New Roman",
                                    "font_size_pt": 16,
                                    "is_bold": True,
                                },
                            },
                            {
                                "id": "p_0_r_0__seg_2",
                                "sourceRunId": "p_0_r_0",
                                "text": "报告",
                                "style": {
                                    "font_name": "FangSong_GB2312",
                                    "font_size_pt": 16,
                                    "is_bold": True,
                                },
                            },
                        ],
                        "metadata": {
                            "inputDocxPath": str(source),
                            "structureIndex": {
                                "paragraphs": [
                                    {
                                        "id": "p_0",
                                        "runNodeIds": [
                                            "p_0_r_0__seg_0",
                                            "p_0_r_0__seg_1",
                                            "p_0_r_0__seg_2",
                                        ],
                                    }
                                ]
                            },
                        },
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            write_docx_from_ir(payload_path, output)

            result = docx.Document(output)
            paragraph = result.paragraphs[0]
            self.assertEqual([run.text for run in paragraph.runs], ["这是", "NLP2025", "报告"])
            self.assertEqual([run.font.name for run in paragraph.runs], ["FangSong_GB2312", "Times New Roman", "FangSong_GB2312"])
            self.assertEqual([run.bold for run in paragraph.runs], [True, True, True])

    @staticmethod
    def _import_docx():
        try:
            import docx  # type: ignore
        except Exception as exc:  # pragma: no cover
            raise unittest.SkipTest(f"python-docx not available: {exc}") from exc
        return docx

    @staticmethod
    def _import_writer():
        from scripts.write_docx_from_ir import write_docx_from_ir  # type: ignore
        return write_docx_from_ir


if __name__ == "__main__":
    unittest.main()
