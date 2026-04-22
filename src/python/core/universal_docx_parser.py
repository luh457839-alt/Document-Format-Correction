from __future__ import annotations

import logging
import re
import warnings
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any, Iterator

from .project_paths import AGENT_MEDIA_DIR

try:
    from lxml import etree

    _LXML_IMPORT_ERROR: Exception | None = None
except Exception as exc:  # pragma: no cover - environment dependent
    etree = None  # type: ignore[assignment]
    _LXML_IMPORT_ERROR = exc

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

try:
    from docx import Document
    from docx.document import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    from docx.text.run import Run

    _DOCX_IMPORT_ERROR: Exception | None = None
except Exception as exc:  # pragma: no cover - environment dependent
    Document = None  # type: ignore[assignment]
    DocxDocument = None  # type: ignore[assignment]
    qn = None  # type: ignore[assignment]
    _Cell = None  # type: ignore[assignment]
    Table = None  # type: ignore[assignment]
    Paragraph = None  # type: ignore[assignment]
    Run = None  # type: ignore[assignment]
    _DOCX_IMPORT_ERROR = exc


LOGGER = logging.getLogger(__name__)

NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


class UnrecognizedNodeWarning(Warning):
    pass


@dataclass
class _DocDefaults:
    font_name: str = "Times New Roman"
    font_size_pt: float = 12.0
    font_color: str = "000000"
    is_bold: bool = False
    is_italic: bool = False


class UniversalDocxParser:
    """
    Parse .docx into a normalized JSON-like state tree for agent observation.
    """

    _OMML_TO_LATEX_XSLT_RAW = (
        r"""
<xsl:stylesheet version="1.0"
    xmlns:xsl="http://www.w3.org/1999/XSL/Transform"
    xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
  <xsl:output method="text" encoding="UTF-8"/>

  <xsl:template match="/">
    <xsl:apply-templates/>
  </xsl:template>

  <xsl:template match="m:oMathPara">
    <xsl:apply-templates select="m:oMath"/>
  </xsl:template>

  <xsl:template match="m:oMath">
    <xsl:for-each select="*">
      <xsl:apply-templates select="."/>
    </xsl:for-each>
  </xsl:template>

  <xsl:template match="m:r">
    <xsl:apply-templates/>
  </xsl:template>

  <xsl:template match="m:t">
    <xsl:value-of select="."/>
  </xsl:template>

  <xsl:template match="m:f">
    <xsl:text>\frac{</xsl:text>
    <xsl:apply-templates select="m:num"/>
    <xsl:text>}{</xsl:text>
    <xsl:apply-templates select="m:den"/>
    <xsl:text>}</xsl:text>
  </xsl:template>

  <xsl:template match="m:num|m:den|m:e|m:sup|m:sub|m:deg">
    <xsl:for-each select="*">
      <xsl:apply-templates select="."/>
    </xsl:for-each>
  </xsl:template>

  <xsl:template match="m:sSup">
    <xsl:apply-templates select="m:e"/>
    <xsl:text>^{</xsl:text>
    <xsl:apply-templates select="m:sup"/>
    <xsl:text>}</xsl:text>
  </xsl:template>

  <xsl:template match="m:sSub">
    <xsl:apply-templates select="m:e"/>
    <xsl:text>_{</xsl:text>
    <xsl:apply-templates select="m:sub"/>
    <xsl:text>}</xsl:text>
  </xsl:template>

  <xsl:template match="m:sSubSup">
    <xsl:apply-templates select="m:e"/>
    <xsl:text>_{</xsl:text>
    <xsl:apply-templates select="m:sub"/>
    <xsl:text>}^{</xsl:text>
    <xsl:apply-templates select="m:sup"/>
    <xsl:text>}</xsl:text>
  </xsl:template>

  <xsl:template match="m:rad">
    <xsl:choose>
      <xsl:when test="m:deg/*">
        <xsl:text>\sqrt[</xsl:text>
        <xsl:apply-templates select="m:deg"/>
        <xsl:text>]{</xsl:text>
        <xsl:apply-templates select="m:e"/>
        <xsl:text>}</xsl:text>
      </xsl:when>
      <xsl:otherwise>
        <xsl:text>\sqrt{</xsl:text>
        <xsl:apply-templates select="m:e"/>
        <xsl:text>}</xsl:text>
      </xsl:otherwise>
    </xsl:choose>
  </xsl:template>

  <xsl:template match="m:d">
    <xsl:text>\left(</xsl:text>
    <xsl:apply-templates select="m:e"/>
    <xsl:text>\right)</xsl:text>
  </xsl:template>

  <xsl:template match="m:nary">
    <xsl:text>\sum</xsl:text>
    <xsl:if test="m:sub/*">
      <xsl:text>_{</xsl:text>
      <xsl:apply-templates select="m:sub"/>
      <xsl:text>}</xsl:text>
    </xsl:if>
    <xsl:if test="m:sup/*">
      <xsl:text>^{</xsl:text>
      <xsl:apply-templates select="m:sup"/>
      <xsl:text>}</xsl:text>
    </xsl:if>
    <xsl:apply-templates select="m:e"/>
  </xsl:template>

  <xsl:template match="m:groupChr|m:limLow|m:limUpp|m:box|m:acc|m:bar|m:sPre">
    <xsl:apply-templates/>
  </xsl:template>

  <xsl:template match="m:*">
    <xsl:apply-templates/>
  </xsl:template>
</xsl:stylesheet>
""".strip()
    )

    def __init__(self, media_dir: str | Path = AGENT_MEDIA_DIR) -> None:
        if etree is None:
            raise RuntimeError(f"lxml is required for UniversalDocxParser: {_LXML_IMPORT_ERROR}")
        self.media_dir = Path(media_dir).resolve()
        self.media_dir.mkdir(parents=True, exist_ok=True)
        self._omml_transform = etree.XSLT(etree.XML(self._OMML_TO_LATEX_XSLT_RAW))

        self._paragraph_idx = 0
        self._table_idx = 0
        self._image_idx = 0
        self._formula_idx = 0
        self._paragraph_records: list[dict[str, Any]] = []

    def parse(self, docx_path: str | Path) -> dict[str, Any]:
        if Document is None or qn is None:
            raise RuntimeError(
                f"python-docx is required for UniversalDocxParser: {_DOCX_IMPORT_ERROR}"
            )
        document = Document(str(docx_path))
        doc_defaults = self._extract_doc_defaults(document)

        self._paragraph_idx = 0
        self._table_idx = 0
        self._image_idx = 0
        self._formula_idx = 0
        self._paragraph_records = []

        nodes: list[dict[str, Any]] = []
        for block in self._iter_block_items(document):
            if isinstance(block, Paragraph):
                nodes.append(self._parse_paragraph_node(block, doc_defaults))
            elif isinstance(block, Table):
                nodes.append(self._parse_table_node(block, doc_defaults))
            else:
                self._warn_unrecognized(f"Unknown top-level block type: {type(block)!r}")

        return {
            "document_meta": {
                "total_paragraphs": len(document.paragraphs),
                "total_tables": len(document.tables),
            },
            "paragraphs": self._paragraph_records,
            "nodes": nodes,
        }

    def _parse_paragraph_node(
        self, paragraph: Paragraph, defaults: _DocDefaults
    ) -> dict[str, Any]:
        paragraph_id = f"p_{self._paragraph_idx}"
        paragraph_index = self._paragraph_idx
        self._paragraph_idx += 1

        children: list[dict[str, Any]] = []
        for run_idx, run in enumerate(paragraph.runs):
            children.extend(
                self._parse_run_children(
                    run=run,
                    text_id=f"p_{paragraph_index}_r_{run_idx}",
                    style=self._resolve_run_style(run, paragraph, defaults),
                )
            )

        self._paragraph_records.append(
            self._build_paragraph_record(paragraph, paragraph_id, children, in_table=False)
        )
        return {"id": paragraph_id, "node_type": "paragraph", "children": children}

    def _parse_table_node(self, table: Table, defaults: _DocDefaults) -> dict[str, Any]:
        table_index = self._table_idx
        table_id = f"tbl_{table_index}"
        self._table_idx += 1

        rows: list[dict[str, Any]] = []
        for row_idx, row in enumerate(table.rows):
            row_data = {"row_index": row_idx, "cells": []}
            for col_idx, cell in enumerate(row.cells):
                cell_data: dict[str, Any] = {
                    "cell_index": col_idx,
                    "paragraphs": [],
                    "tables": [],
                }

                paragraph_local_idx = 0
                for block in self._iter_block_items(cell):
                    if isinstance(block, Paragraph):
                        paragraph_data = {"node_type": "paragraph", "children": []}
                        for run_idx, run in enumerate(block.runs):
                            text_id = (
                                f"tbl_{table_index}_r_{row_idx}_c_{col_idx}"
                                f"_p_{paragraph_local_idx}_r_{run_idx}"
                            )
                            paragraph_data["children"].extend(
                                self._parse_run_children(
                                    run=run,
                                    text_id=text_id,
                                    style=self._resolve_run_style(run, block, defaults),
                                )
                            )
                        self._paragraph_records.append(
                            self._build_paragraph_record(
                                block,
                                f"tbl_{table_index}_r_{row_idx}_c_{col_idx}_p_{paragraph_local_idx}",
                                paragraph_data["children"],
                                in_table=True,
                            )
                        )
                        cell_data["paragraphs"].append(paragraph_data)
                        paragraph_local_idx += 1
                    elif isinstance(block, Table):
                        cell_data["tables"].append(self._parse_table_node(block, defaults))
                    else:
                        self._warn_unrecognized(
                            f"Unknown cell block type: {type(block)!r} "
                            f"(table={table_index}, row={row_idx}, col={col_idx})"
                        )

                row_data["cells"].append(cell_data)
            rows.append(row_data)

        return {"id": table_id, "node_type": "table", "rows": rows}

    def _parse_run_children(
        self, run: Run, text_id: str, style: dict[str, Any]
    ) -> list[dict[str, Any]]:
        children: list[dict[str, Any]] = []
        text = run.text or ""
        if text:
            children.append(
                {
                    "id": text_id,
                    "node_type": "text_run",
                    "content": text,
                    "style": style,
                }
            )

        for child in run._element:
            local_name = etree.QName(child).localname
            if local_name == "drawing":
                children.append(self._extract_image_node(run, child))
                continue
            if local_name == "pict":
                image_nodes = self._extract_vml_images(run, child)
                children.extend(image_nodes)
                continue
            if child.tag in {qn("m:oMath"), qn("m:oMathPara")}:
                children.append(self._extract_formula_node(child))
                continue

            if local_name not in {"t", "rPr", "br", "tab", "cr", "noBreakHyphen"}:
                self._warn_unrecognized(f"Unknown run child tag skipped: {child.tag}")

        return children

    def _extract_image_node(self, run: Run, drawing_element: etree._Element) -> dict[str, Any]:
        img_id = f"img_{self._image_idx}"
        self._image_idx += 1

        rid = None
        blip = self._find_first_descendant(drawing_element, qn("a:blip"))
        if blip is not None:
            rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))

        size = self._extract_drawing_size(drawing_element)
        if rid is None:
            self._warn_unrecognized(f"Drawing image without relationship id, node={img_id}")
            return {"id": img_id, "node_type": "image", "src": "extraction_failed", "size": size}

        return self._export_image_by_rid(run, rid, img_id, size)

    def _extract_vml_images(self, run: Run, pict_element: etree._Element) -> list[dict[str, Any]]:
        nodes: list[dict[str, Any]] = []
        image_data_tags = self._find_descendants(pict_element, self._ns_tag("v", "imagedata"))
        for image_data in image_data_tags:
            img_id = f"img_{self._image_idx}"
            self._image_idx += 1
            rid = image_data.get(qn("r:id"))
            size = self._extract_vml_size(image_data.getparent())
            if rid is None:
                self._warn_unrecognized(f"VML image without relationship id, node={img_id}")
                nodes.append(
                    {"id": img_id, "node_type": "image", "src": "extraction_failed", "size": size}
                )
                continue
            nodes.append(self._export_image_by_rid(run, rid, img_id, size))
        return nodes

    def _export_image_by_rid(
        self, run: Run, rid: str, image_id: str, size: dict[str, float]
    ) -> dict[str, Any]:
        try:
            part = run.part.related_parts[rid]
            suffix = Path(part.partname).suffix or ".bin"
            output_path = (self.media_dir / f"{image_id}{suffix}").resolve()
            output_path.write_bytes(part.blob)
            return {
                "id": image_id,
                "node_type": "image",
                "src": str(output_path),
                "size": size,
            }
        except Exception as exc:  # pragma: no cover - depends on damaged package data
            LOGGER.warning("Image extraction failed for %s: %s", image_id, exc)
            return {
                "id": image_id,
                "node_type": "image",
                "src": "extraction_failed",
                "size": size,
            }

    def _extract_formula_node(self, math_element: etree._Element) -> dict[str, Any]:
        formula_id = f"math_{self._formula_idx}"
        self._formula_idx += 1
        latex = self._omml_to_latex(math_element)
        return {
            "id": formula_id,
            "node_type": "formula",
            "format": "latex",
            "content": latex,
        }

    def _omml_to_latex(self, math_element: etree._Element) -> str:
        try:
            tree = etree.ElementTree(math_element)
            transformed = self._omml_transform(tree)
            latex = str(transformed).strip()
            return latex if latex else "\\text{unparsed_formula}"
        except Exception as exc:  # pragma: no cover - depends on malformed OMML
            LOGGER.warning("OMML conversion failed: %s", exc)
            self._warn_unrecognized(f"OMML conversion failed: {exc}")
            return "\\text{omml_conversion_failed}"

    def _extract_doc_defaults(self, document: DocxDocument) -> _DocDefaults:
        defaults = _DocDefaults()
        styles_el = document.styles.element

        fonts = self._find_first_by_path(
            styles_el,
            [qn("w:docDefaults"), qn("w:rPrDefault"), qn("w:rPr"), qn("w:rFonts")],
        )
        if fonts is not None:
            default_font = (
                fonts.get(qn("w:eastAsia"))
                or fonts.get(qn("w:ascii"))
                or fonts.get(qn("w:hAnsi"))
                or fonts.get(qn("w:cs"))
            )
            if default_font:
                defaults.font_name = default_font

        sz = self._find_first_by_path(
            styles_el,
            [qn("w:docDefaults"), qn("w:rPrDefault"), qn("w:rPr"), qn("w:sz")],
        )
        if sz is not None and sz.get(qn("w:val")):
            try:
                defaults.font_size_pt = int(sz.get(qn("w:val"))) / 2.0
            except ValueError:
                pass

        color = self._find_first_by_path(
            styles_el,
            [qn("w:docDefaults"), qn("w:rPrDefault"), qn("w:rPr"), qn("w:color")],
        )
        if color is not None and color.get(qn("w:val")) and color.get(qn("w:val")) != "auto":
            defaults.font_color = color.get(qn("w:val")).upper()

        return defaults

    def _resolve_run_style(
        self, run: Run, paragraph: Paragraph, defaults: _DocDefaults
    ) -> dict[str, Any]:
        font_name = self._first_non_none(
            run.font.name,
            getattr(getattr(run, "style", None), "font", None)
            and getattr(run.style.font, "name", None),
            getattr(paragraph.style.font, "name", None),
            self._normal_style_attr(paragraph, "name"),
            defaults.font_name,
        )
        font_size_pt = self._first_non_none(
            run.font.size.pt if run.font.size else None,
            (
                getattr(run.style.font, "size", None).pt
                if getattr(getattr(run, "style", None), "font", None)
                and getattr(run.style.font, "size", None)
                else None
            ),
            paragraph.style.font.size.pt if paragraph.style.font.size else None,
            self._normal_style_attr(paragraph, "size"),
            defaults.font_size_pt,
        )

        font_color = self._first_non_none(
            self._rgb_to_hex(run.font.color.rgb if run.font.color else None),
            (
                self._rgb_to_hex(run.style.font.color.rgb)
                if getattr(getattr(run, "style", None), "font", None)
                and getattr(run.style.font, "color", None)
                else None
            ),
            self._rgb_to_hex(paragraph.style.font.color.rgb if paragraph.style.font.color else None),
            self._normal_style_attr(paragraph, "color"),
            defaults.font_color,
        )

        is_bold = bool(
            self._first_non_none(
                run.bold,
                getattr(getattr(run, "style", None), "font", None)
                and getattr(run.style.font, "bold", None),
                paragraph.style.font.bold,
                self._normal_style_attr(paragraph, "bold"),
                defaults.is_bold,
            )
        )
        is_italic = bool(
            self._first_non_none(
                run.italic,
                getattr(getattr(run, "style", None), "font", None)
                and getattr(run.style.font, "italic", None),
                paragraph.style.font.italic,
                self._normal_style_attr(paragraph, "italic"),
                defaults.is_italic,
            )
        )

        style = {
            "font_name": str(font_name or defaults.font_name),
            "font_size_pt": float(font_size_pt or defaults.font_size_pt),
            "font_color": str(font_color or defaults.font_color).upper(),
            "is_bold": is_bold,
            "is_italic": is_italic,
        }
        explicit_line_spacing = self._read_explicit_line_spacing(paragraph)
        if explicit_line_spacing is not None:
            style["line_spacing"] = explicit_line_spacing
        return style

    def _read_explicit_line_spacing(self, paragraph: Paragraph) -> float | dict[str, float] | None:
        paragraph_element = getattr(paragraph, "_p", None)
        if paragraph_element is None:
            return None
        p_pr = self._find_first_child(paragraph_element, qn("w:pPr"))
        if p_pr is None:
            return None
        spacing = self._find_first_child(p_pr, qn("w:spacing"))
        if spacing is None:
            return None
        line_val = spacing.get(qn("w:line"))
        if line_val is None:
            return None
        try:
            line = int(line_val)
        except ValueError:
            return None
        if line <= 0:
            return None
        line_rule = str(spacing.get(qn("w:lineRule")) or "auto").strip().lower()
        if line_rule == "exact":
            return {"mode": "exact", "pt": round(line / 20.0, 4)}
        if line_rule in {"", "auto"}:
            return round(line / 240.0, 4)
        return None

    def _build_paragraph_record(
        self,
        paragraph: Paragraph,
        paragraph_id: str,
        children: list[dict[str, Any]],
        *,
        in_table: bool,
    ) -> dict[str, Any]:
        role_info = self._classify_paragraph_role(paragraph, in_table=in_table)
        record: dict[str, Any] = {
            "id": paragraph_id,
            "text": "".join(
                str(child.get("content", ""))
                for child in children
                if isinstance(child, dict) and child.get("node_type") == "text_run"
            ).strip(),
            "role": role_info["role"],
            "style_name": self._paragraph_style_name(paragraph),
            "run_ids": [
                str(child["id"]).strip()
                for child in children
                if isinstance(child, dict)
                and child.get("node_type") == "text_run"
                and isinstance(child.get("id"), str)
                and str(child["id"]).strip()
            ],
            "in_table": in_table,
        }
        if role_info.get("heading_level") is not None:
            record["heading_level"] = role_info["heading_level"]
        if role_info.get("list_level") is not None:
            record["list_level"] = role_info["list_level"]
        return record

    def _classify_paragraph_role(self, paragraph: Paragraph, *, in_table: bool) -> dict[str, Any]:
        style_name = self._paragraph_style_name(paragraph)
        heading_level = self._extract_heading_level(style_name)
        if heading_level is not None:
            return {"role": "heading", "heading_level": heading_level}
        compact_style = re.sub(r"[\s_-]+", "", style_name).lower()
        if compact_style in {"title", "subtitle"}:
            return {"role": "title"}
        list_level = self._extract_list_level(paragraph)
        if list_level is not None:
            return {"role": "list_item", "list_level": list_level}
        if in_table:
            return {"role": "table_text"}
        if str(getattr(paragraph, "text", "") or "").strip():
            return {"role": "body"}
        return {"role": "unknown"}

    def _paragraph_style_name(self, paragraph: Paragraph) -> str:
        style_name = getattr(getattr(paragraph, "style", None), "name", "")
        return str(style_name or "").strip()

    def _extract_heading_level(self, style_name: str) -> int | None:
        compact_style = re.sub(r"[\s_-]+", "", style_name).lower()
        match = re.match(r"heading(\d+)$", compact_style)
        if match:
            return int(match.group(1))
        match = re.match(r"标题(\d+)$", style_name.replace(" ", ""))
        if match:
            return int(match.group(1))
        return None

    def _extract_list_level(self, paragraph: Paragraph) -> int | None:
        paragraph_element = getattr(paragraph, "_p", None)
        if paragraph_element is None:
            return None
        fake_list_level = getattr(paragraph_element, "list_level", None)
        if isinstance(fake_list_level, int):
            return fake_list_level
        if not callable(getattr(paragraph_element, "iterchildren", None)) and not hasattr(paragraph_element, "__iter__"):
            return None
        p_pr = self._find_first_child(paragraph_element, qn("w:pPr"))
        if p_pr is None:
            return None
        num_pr = self._find_first_child(p_pr, qn("w:numPr"))
        if num_pr is None:
            return None
        ilvl = self._find_first_child(num_pr, qn("w:ilvl"))
        if ilvl is None:
            return 0
        value = ilvl.get(qn("w:val"))
        if value is None:
            return 0
        try:
            return int(value)
        except ValueError:
            return 0

    def _normal_style_attr(self, paragraph: Paragraph, attr: str) -> Any:
        normal = paragraph.part.document.styles["Normal"].font
        if attr == "size":
            return normal.size.pt if normal.size else None
        if attr == "color":
            return self._rgb_to_hex(normal.color.rgb if normal.color else None)
        return getattr(normal, attr, None)

    def _extract_drawing_size(self, drawing_element: etree._Element) -> dict[str, float]:
        extent = self._find_first_descendant(drawing_element, qn("wp:extent"))
        if extent is None:
            return {"width": 0.0, "height": 0.0}
        cx = float(extent.get("cx", 0.0))
        cy = float(extent.get("cy", 0.0))
        # EMU -> px, at 96 dpi.
        return {"width": round(cx / 9525.0, 2), "height": round(cy / 9525.0, 2)}

    def _extract_vml_size(self, shape_element: etree._Element | None) -> dict[str, float]:
        if shape_element is None:
            return {"width": 0.0, "height": 0.0}
        style_attr = shape_element.get("style") or ""
        width = self._parse_vml_dimension(style_attr, "width")
        height = self._parse_vml_dimension(style_attr, "height")
        return {"width": width, "height": height}

    def _parse_vml_dimension(self, style_attr: str, prop_name: str) -> float:
        pattern = rf"{prop_name}\s*:\s*([0-9.]+)(pt|px)?"
        match = re.search(pattern, style_attr, re.IGNORECASE)
        if not match:
            return 0.0
        value = float(match.group(1))
        unit = (match.group(2) or "px").lower()
        if unit == "pt":
            return round(value * (96.0 / 72.0), 2)
        return round(value, 2)

    def _iter_block_items(self, parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
        if DocxDocument is not None and isinstance(parent, DocxDocument):
            parent_element = parent.element.body
            parent_obj: Any = parent
        elif _Cell is not None and isinstance(parent, _Cell):
            parent_element = parent._tc
            parent_obj = parent
        else:
            raise TypeError(f"Unsupported parent type for block iteration: {type(parent)!r}")

        for child in parent_element.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent_obj)
                continue
            if child.tag == qn("w:tbl"):
                yield Table(child, parent_obj)
                continue
            if child.tag == qn("w:sectPr"):
                continue
            self._warn_unrecognized(f"Unknown block tag skipped: {child.tag}")

    def _warn_unrecognized(self, message: str) -> None:
        warnings.warn(message, UnrecognizedNodeWarning, stacklevel=2)
        LOGGER.warning(message)

    def _find_first_by_path(self, element: Any, tags: list[str]) -> Any | None:
        current = element
        for tag in tags:
            current = self._find_first_child(current, tag)
            if current is None:
                return None
        return current

    def _find_first_child(self, element: Any, tag: str) -> Any | None:
        for child in self._iter_children(element):
            if getattr(child, "tag", None) == tag:
                return child
        return None

    def _find_first_descendant(self, element: Any, tag: str) -> Any | None:
        for child in self._iter_descendants(element):
            if getattr(child, "tag", None) == tag:
                return child
        return None

    def _find_descendants(self, element: Any, tag: str) -> list[Any]:
        return [child for child in self._iter_descendants(element) if getattr(child, "tag", None) == tag]

    def _iter_descendants(self, element: Any) -> Iterator[Any]:
        for child in self._iter_children(element):
            yield child
            yield from self._iter_descendants(child)

    @staticmethod
    def _iter_children(element: Any) -> Iterator[Any]:
        iterchildren = getattr(element, "iterchildren", None)
        if callable(iterchildren):
            yield from iterchildren()
            return
        yield from iter(element)

    @staticmethod
    def _ns_tag(prefix: str, local_name: str) -> str:
        return f"{{{NS[prefix]}}}{local_name}"

    @staticmethod
    def _rgb_to_hex(rgb: Any) -> str | None:
        if rgb is None:
            return None
        value = str(rgb)
        if not value:
            return None
        return value.upper().replace("#", "")

    @staticmethod
    def _first_non_none(*values: Any) -> Any:
        for value in values:
            if value is not None:
                return value
        return None


__all__ = ["UniversalDocxParser", "UnrecognizedNodeWarning"]
